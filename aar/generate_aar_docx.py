import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print('python-docx not installed. Run: pip install python-docx')
    sys.exit(1)

base_dir = Path(r"C:/Users/Mizzi/APEX_NEXUS_SYSTEM/04_PROJECTS/DIDGERI_BOOM/aar")
content_file = base_dir / "aar_content.txt"
output_file = base_dir / "CON277_AAR.docx"

if not content_file.exists():
    print(f"Content file not found: {content_file}")
    sys.exit(1)

with open(content_file, "r", encoding="utf-8") as f:
    lines = [line.strip() for line in f if line.strip()]

doc = Document()
doc.add_heading("After Action Review - CON 277", level=1)

for line in lines:
    if line.startswith("## "):
        heading = line[3:].strip()
        doc.add_heading(heading, level=2)
    elif line.startswith("# "):
        pass
    elif line.startswith("- "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
    else:
        doc.add_paragraph(line)

doc.save(str(output_file))
print(f"AAR document generated at {output_file}")
